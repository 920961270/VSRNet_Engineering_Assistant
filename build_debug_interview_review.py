"""Build the debug and interview review Markdown/DOCX documents."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCS_DIR = ROOT / "docs"
METRICS_PATH = ROOT / "outputs" / "metrics.json"
REPORT_PATH = ROOT / "outputs" / "engineering_report.md"
MD_PATH = DOCS_DIR / "VSRNet_Engineering_Assistant_Debug_and_Interview_Review.md"
DOCX_PATH = DOCS_DIR / "VSRNet_Engineering_Assistant_Debug_and_Interview_Review.docx"


def load_metrics() -> dict:
    """Load latest metrics if available."""
    if not METRICS_PATH.exists():
        return {}
    return json.loads(METRICS_PATH.read_text(encoding="utf-8"))


def fmt(metrics: dict, key: str, digits: int = 3, default: str = "取决于最新一次运行") -> str:
    """Format a metric value."""
    value = metrics.get(key)
    if value is None:
        return default
    if isinstance(value, float):
        return f"{value:.{digits}f}".rstrip("0").rstrip(".")
    return str(value)


def metric_summary(metrics: dict) -> str:
    """Return a compact metric interpretation paragraph."""
    if not metrics:
        return "当前没有读取到 outputs/metrics.json，具体指标取决于最新一次运行。"
    return (
        f"最近一次 metrics.json 显示：PSNR={fmt(metrics, 'psnr')} dB，"
        f"SSIM={fmt(metrics, 'ssim')}，jitter={fmt(metrics, 'jitter')} 像素，"
        f"IoU continuity={fmt(metrics, 'iou_continuity')}，"
        f"confidence_mean={fmt(metrics, 'confidence_mean')}，"
        f"detection_count={fmt(metrics, 'detection_count', 0)}，"
        f"detected_frame_count={fmt(metrics, 'detected_frame_count', 0)}，"
        f"restoration_total_time_sec={fmt(metrics, 'restoration_total_time_sec')} 秒，"
        f"inference_time_ms={fmt(metrics, 'inference_time_ms')} ms/frame。"
    )


def build_markdown(metrics: dict) -> str:
    """Build the full Markdown review."""
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    metrics_text = metric_summary(metrics)
    psnr = fmt(metrics, "psnr")
    ssim = fmt(metrics, "ssim")
    jitter = fmt(metrics, "jitter")
    iou = fmt(metrics, "iou_continuity")
    conf_mean = fmt(metrics, "confidence_mean")
    conf_var = fmt(metrics, "confidence_variance")
    det_count = fmt(metrics, "detection_count", 0)
    det_frames = fmt(metrics, "detected_frame_count", 0)
    frame_count = fmt(metrics, "restored_frame_count", 0)
    total_time = fmt(metrics, "restoration_total_time_sec")
    infer_ms = fmt(metrics, "inference_time_ms")
    duration = fmt(metrics, "restored_duration_sec")
    fps = fmt(metrics, "fps")
    warnings = metrics.get("warnings", "取决于最新一次运行")
    if warnings == []:
        warnings = "无明确 warning"

    return f"""# VSRNet Engineering Assistant 调试与面试复盘文档

生成时间：{generated_at}  
项目目录：`G:\\4周学习计划\\第三周\\VSRNet_Engineering_Assistant`

## 项目一句话总结

这是一个面向振动退化监控视频的 CV + RAG + LLM 工程评估助手：它复用已有 VSRNet 恢复模型，接入 YOLO 下游检测，计算图像质量与检测稳定性指标，再用 FAISS RAG 检索项目知识，最后生成可解释、可复用、带防幻觉机制的工程分析报告。

## 1. 项目背景与目标

### 1.1 为什么要做 VSRNet Engineering Assistant

单纯训练或调用一个视频恢复模型，并不能直接回答“这个模型在工程上是否有价值”。VSRNet 的目标场景是振动退化监控视频，真实工程问题不是只看画面是否变清晰，还要看恢复后的视频是否有助于下游检测、跟踪和稳定分析。因此，本项目把视频恢复、YOLO 检测、指标计算、项目知识检索和报告生成串成一个小型 AI 工程评估系统。

### 1.2 项目想解决的问题：从视频恢复到工程分析

这个项目把原始问题拆成四层：

- 第一层：给定振动模糊视频，尽可能调用已有 VSRNet 代码完成恢复。
- 第二层：用 YOLO 在恢复结果上检测目标，观察下游感知是否稳定。
- 第三层：计算 PSNR、SSIM、jitter、IoU continuity、confidence mean、inference time 等指标。
- 第四层：结合 RAG 检索到的 VSRNet 项目知识，生成可解释的工程报告。

### 1.3 和求职 JD 的关系：CV + RAG + LLM + 工程系统

这个项目可以对应多类实习 JD：计算机视觉、视频增强、多模态/LLM 应用、RAG、AI 工程系统、MLOps 雏形。它不是只写一个 demo，而是展示了“把模型接进工程流程，并能解释结果”的能力。

### 1.4 最终项目定位：不是单一模型，而是小型 AI 工程评估助手

项目定位应表达为：我没有重新发明 VSRNet，而是把已有模型包装成可运行的评估流程，并在其后接入检测、指标、RAG 和报告。这种定位比“我训练了一个模型”更适合工程实习面试。

## 2. 系统整体流程

### 2.1 Stage 1：Placeholder MVP 骨架

Stage 1 先搭建骨架：`main_pipeline.py`、`modules/restoration.py`、`modules/detection.py`、`modules/metrics.py`、`modules/rag_retriever.py`、`modules/report_generator.py`。当真实 VSRNet 或 YOLO 未接入时，placeholder 保证流程不断裂。这不是错误，而是工程上常见的“先跑通接口，再替换重组件”的做法。

### 2.2 Stage 2：真实 VSRNet + YOLO + Metrics

Stage 2 开始调用 `G:\\EDVR_3levels\\VSRNet\\infer_vsrnet_video.py`，并把恢复结果传给 YOLO。随后 `modules/metrics.py` 计算恢复质量和检测稳定性指标。最近一次运行的核心指标为：{metrics_text}

### 2.3 RAG 检索模块：项目知识库增强报告

`modules/rag_retriever.py` 使用 sentence-transformers + FAISS `IndexFlatIP`，对 `documents/` 下的 VSRNet 文本进行向量检索。RAG 的意义是让报告引用项目本身的背景、方法和评估说明，而不是只凭模板或 LLM 自由发挥。

### 2.4 Rule-Based Engineering Analysis：稳定可信的指标解释

Rule-based analysis 是当前报告里的可信主分析。它直接由 metrics 字典生成，逻辑确定、可复现，不会像 LLM 那样夸大“near-perfect”“real-time”。例如 PSNR={psnr}、SSIM={ssim} 时，系统会解释为“moderate restoration quality”，而不是“near-perfect restoration”。

### 2.5 LLM Additional Analysis：可选的补充评论层

LLM 只做附加评论，不替代 rule-based 分析。`modules/llm_report.py` 会保存 raw output、sanitize 后的 output 和 validation log。如果 LLM 编造数字、过度夸大、或与 rule-based 结论冲突，就不会作为可靠结论写入。

### 2.6 Report-only / Reuse 模式：避免重复推理

VSRNet 推理耗时较长，最近一次约 {total_time} 秒、{infer_ms} ms/frame。因此项目加入 `--skip-restoration`、`--skip-detection`、`--skip-metrics`，可以复用已有 restored video、detection CSV 和 metrics JSON，只重新做 RAG 与报告生成。

## 3. 当前项目目录结构与文件作用

| 文件/目录 | 作用 |
| --- | --- |
| `main_pipeline.py` | 命令行入口，串联 restoration、detection、metrics、RAG、rule-based analysis、LLM additional analysis 和报告生成。 |
| `modules/restoration.py` | VSRNet wrapper。负责检查输入视频、checkpoint、调用外部推理脚本、处理绝对路径、记录推理耗时。 |
| `modules/detection.py` | YOLO wrapper。优先使用 Ultralytics YOLO 检测/跟踪，输出 `detection_results.csv`。 |
| `modules/metrics.py` | 计算 PSNR、SSIM、jitter、IoU continuity、confidence、FPS、duration、inference time 等指标。 |
| `modules/rag_retriever.py` | 使用本地 all-MiniLM-L6-v2 + FAISS 检索项目知识文档。 |
| `modules/llm_report.py` | 生成 rule-based 主分析，并可选调用 OpenAI/Ollama 生成 LLM Additional Analysis，同时做 sanitize 与 validation。 |
| `modules/report_generator.py` | 把指标、RAG 结果、rule-based 分析、LLM 补充评论写入 `engineering_report.md`。 |
| `documents/` | RAG 知识库，存放 VSRNet summary、method notes、evaluation notes、experiment logs。 |
| `inputs/` | 输入视频和 GT 视频。 |
| `outputs/` | 输出目录，包含恢复视频、检测 CSV、指标和最终报告。 |
| `outputs/metrics.json` | 当前最重要的机器可读指标文件。 |
| `outputs/metrics.csv` | 指标 CSV，方便表格查看。 |
| `outputs/detection_results.csv` | YOLO 检测/跟踪结果，是 jitter 和 IoU continuity 的来源。 |
| `outputs/engineering_report.md` | 最终 Markdown 报告。 |
| `outputs/llm_raw_output.md` | LLM 原始输出，便于调试小模型。 |
| `outputs/llm_sanitized_output.md` | 经过不安全措辞替换后的 LLM 输出。 |
| `outputs/llm_validation_log.txt` | 记录匹配 phrase、句子、动作、最终验证结果。 |

## 4. 运行命令与运行位置

### 4.1 为什么不要在 README 绿色按钮里运行

README 里的代码块只是说明文档，不是交互式终端。多行 PowerShell 命令需要在 PowerShell、PyCharm Terminal 或系统终端里运行。尤其是带 `^` 的 Windows 多行命令，如果复制位置不对，很容易丢失参数，例如漏掉 `--question`。

### 4.2 PowerShell / PyCharm Terminal / Ollama 聊天窗口的区别

- PowerShell：运行 `python main_pipeline.py ...` 的位置。
- PyCharm Terminal：本质也是终端，可以运行同样命令，但要注意当前工作目录。
- Ollama 聊天窗口：用于和模型交互，不适合直接运行 Python pipeline。

### 4.3 Full Run 完整运行命令

```powershell
python main_pipeline.py ^
  --input inputs/high_blur_640_10s.mp4 ^
  --gt inputs/GT_640_10s.mp4 ^
  --question "How does VSRNet perform under high vibration?" ^
  --vsrnet-checkpoint models/vsrnet/best.pth ^
  --yolo-model yolov8n.pt ^
  --documents documents ^
  --output outputs
```

### 4.4 Report-only / LLM-only 复用模式命令

```powershell
python main_pipeline.py ^
  --input inputs/high_blur_640_10s.mp4 ^
  --gt inputs/GT_640_10s.mp4 ^
  --restored-video outputs/high_blur_640_10s_vsrnet_restored.mp4 ^
  --skip-restoration ^
  --detection-csv outputs/detection_results.csv ^
  --skip-detection ^
  --metrics-json outputs/metrics.json ^
  --skip-metrics ^
  --question "How does VSRNet perform under high vibration?" ^
  --documents documents ^
  --output outputs ^
  --use-llm-report ^
  --llm-provider ollama ^
  --llm-model qwen2.5:1.5b
```

### 4.5 Ollama 常用命令：list、pull、run、bye

```powershell
ollama list
ollama pull qwen2.5:1.5b
ollama run qwen2.5:1.5b
/bye
```

如果使用更大的模型，例如 qwen2.5:3b，可以增加 timeout：

```powershell
python main_pipeline.py ^
  --input inputs/high_blur_640_10s.mp4 ^
  --gt inputs/GT_640_10s.mp4 ^
  --restored-video outputs/high_blur_640_10s_vsrnet_restored.mp4 ^
  --skip-restoration ^
  --detection-csv outputs/detection_results.csv ^
  --skip-detection ^
  --metrics-json outputs/metrics.json ^
  --skip-metrics ^
  --question "How does VSRNet perform under high vibration?" ^
  --documents documents ^
  --output outputs ^
  --use-llm-report ^
  --llm-provider ollama ^
  --llm-model qwen2.5:3b ^
  --llm-timeout 300
```

### 4.6 Windows 路径、中文路径与引号注意事项

Windows 路径建议加引号，尤其是包含中文、空格或反斜杠时。项目中曾遇到中文路径少字、目录名不一致导致 checkpoint 找不到的问题。因此命令里应优先使用完整路径或确认当前目录正确。

## 5. 调试问题复盘：从 Stage 1 到 Stage 2

下面每个问题都按“现象 - 原因 - 解决方法 - 面试表达”组织。

### 5.1 `--question` 参数缺失：README 多行命令没有正确执行

- 现象：程序报错提示缺少 required argument `--question`。
- 原因：Windows 多行命令的 `^` 没有正确复制到终端，或在 README 预览区直接运行，导致后续行没有传给 Python。
- 解决方法：在 PowerShell 或 PyCharm Terminal 中运行，并确认每行末尾的 `^` 没有多余空格。
- 面试中可以怎么说：我把命令行参数设计成显式必填，是为了保证报告一定有用户问题；同时通过 README 示例说明 Windows 多行命令的使用方式。

### 5.2 Stage 1 placeholder 是什么，为什么不是报错

- 现象：没有真实 VSRNet 或 YOLO 时，系统仍生成 placeholder 输出。
- 原因：Stage 1 目标是跑通软件骨架，不依赖重模型。
- 解决方法：将 restoration 和 detection 设计成 wrapper，缺组件时返回结构化 placeholder。
- 面试表达：这是工程中的渐进式集成策略，先稳定接口，再替换重组件。

### 5.3 输入视频路径错误：FileNotFoundError

- 现象：VSRNet 推理脚本无法打开输入视频。
- 原因：相对路径、中文路径、当前工作目录不一致。
- 解决方法：wrapper 内部统一转换为 absolute path，并在运行前检查文件存在。
- 面试表达：我没有让底层脚本自己失败，而是在 wrapper 层做了输入校验和可解释错误。

### 5.4 checkpoint 路径拼写错误

- 现象：checkpoint 找不到，真实推理无法启动。
- 原因：中文路径少字、目录名和实际文件不一致。
- 解决方法：`modules/restoration.py` 检查 checkpoint 是否存在；不存在时给出明确 warning。
- 面试表达：模型权重路径是工程系统常见失败点，所以我把路径校验前移。

### 5.5 1080p 视频内存不足：为什么先生成 640 宽 10 秒测试视频

- 现象：高分辨率视频容易显存不足或推理时间过长。
- 原因：视频恢复模型需要多帧输入，显存占用随分辨率和帧数上升。
- 解决方法：先用 640 宽、约 10 秒、250 帧的小样本验证链路。
- 面试表达：我先用可控小样本做端到端验证，再逐步扩大规模，这是降低调试成本的常见方法。

### 5.6 VSRNet 推理成功但主程序找不到输出

- 现象：推理进度 100%，但主程序提示找不到 `outputs/...mp4`。
- 原因：子进程 `cwd` 是 VSRNet repo，传入的相对 `outputs` 路径被解释到了错误目录。
- 解决方法：所有关键路径在 wrapper 内转成绝对路径，包括 input、checkpoint、output video。
- 面试表达：这个 bug 的关键是理解 subprocess 的 working directory 对相对路径的影响。

### 5.7 输出文件夹为空：wrapper 和推理脚本参数没有对齐

- 现象：输出目录里只有空文件夹，没有可用 mp4。
- 原因：wrapper 假设有 frame directory 参数，但实际 `infer_vsrnet_video.py` 主要支持 `--save-video`。
- 解决方法：检查真实 CLI，按脚本实际参数传 `--input-video`、`--ckpt`、`--save-video`。
- 面试表达：我先阅读被包装脚本的真实接口，而不是凭想象写 wrapper。

### 5.8 SSIM 计算失败：彩色图像通道维度与 win_size 问题

- 现象：SSIM 报 `win_size exceeds image extent` 或 multichannel/channel_axis 兼容问题。
- 原因：不同 scikit-image 版本对彩色图像参数不同，直接对 3 通道帧算 SSIM 容易踩坑。
- 解决方法：将 GT 和 restored 帧转成灰度 2D 图，再选择安全 odd `win_size`。
- 面试表达：我把问题从版本兼容转化成更稳定的灰度 SSIM 计算，降低依赖差异风险。

### 5.9 指标 unavailable：哪些是真缺失，哪些只是没实现

- 现象：早期报告中 jitter、IoU continuity、inference time 显示 unavailable。
- 原因：不是数据一定缺失，而是指标计算逻辑还没有实现。
- 解决方法：从 detection CSV 解析 bbox，计算中心点位移和连续帧 IoU；从 VSRNet 子进程计时计算 ms/frame。
- 面试表达：我区分“数据不可得”和“工程逻辑未实现”，逐步补齐指标。

### 5.10 显存不足导致 VSRNet fallback placeholder

- 现象：真实推理失败后系统回退 placeholder。
- 原因：模型或视频规模超过当前机器资源。
- 解决方法：保留 fallback，但在 full run 和 skip mode 中区分：正常模式可 fallback，显式 skip 模式必须严格验证复用文件。
- 面试表达：这体现了 robustness：实验流程不中断，同时显式复用模式不隐藏错误。

### 5.11 为什么要加入 reuse/report-only 模式

- 现象：每次改报告或 LLM prompt 都重新跑 VSRNet，耗时几分钟并占 GPU。
- 原因：管线没有阶段缓存。
- 解决方法：加入 `--restored-video`、`--detection-csv`、`--metrics-json` 和对应 `--skip-*` 参数。
- 面试表达：我把昂贵阶段和轻量报告阶段解耦，提高迭代效率，也降低显存压力。

### 5.12 Ollama timeout：为什么 qwen2.5:3b 需要更长等待时间

- 现象：1.5B 模型快但输出质量不稳定，3B 模型更慢。
- 原因：本地小模型推理速度取决于 CPU/GPU、上下文长度和模型大小。
- 解决方法：加入 `--llm-timeout`，例如 3B 模型用 `--llm-timeout 300`。
- 面试表达：我没有把 timeout 写死，而是暴露为配置参数，适配不同本地环境。

### 5.13 LLM 输出被 rejected：validation 与 sanitize 的作用

- 现象：LLM 可能说“near-perfect”“real-time”等与指标不符的话。
- 原因：小模型容易套用积极模板，或者误读指标。
- 解决方法：保存 raw output，先 sanitize 不安全措辞，再 validation；安全否定句如“not near-perfect”不再误拒绝。
- 面试表达：我把 LLM 当成补充评论层，并用规则校验防止它覆盖确定性指标。

## 6. 指标解释与工程含义

| 指标 | 最新值 | 工程含义 |
| --- | ---: | --- |
| PSNR | {psnr} dB | 约 19.39 dB 表示恢复质量中等/有限，不应称为接近完美。 |
| SSIM | {ssim} | 约 0.699 表示结构相似度中等，不是高结构保真。 |
| jitter | {jitter} px | 约 1.57 像素说明连续帧中心点位移较低，但要结合目标尺度和检测质量看。 |
| IoU continuity | {iou} | 约 0.92 表示连续帧检测框重叠稳定，是当前最积极的信号。 |
| confidence mean | {conf_mean} | 约 0.28 表明检测置信度偏低，检测可靠性仍需提升。 |
| confidence variance | {conf_var} | 反映检测置信度波动情况。 |
| detection count / detected frame count | {det_count} / {det_frames} | 表示在 {frame_count} 帧视频中，部分帧检测到了目标，不是每帧都有可靠检测。 |
| restoration total time | {total_time} sec | 对约 {duration} 秒视频耗时数百秒，说明当前实现计算开销大。 |
| inference time | {infer_ms} ms/frame | 约 1100 ms/frame，明显不是实时。 |
| FPS / duration | {fps} / {duration} sec | 用于判断帧数、视频时长和 per-frame 指标。 |
| warnings | {warnings} | warning 为空时说明本次指标计算链路没有显式异常，但仍需看指标本身。 |

## 7. LLM 模块设计复盘

### 7.1 为什么要加入 LLM：求职与 JD 匹配

LLM 模块让项目不仅是 CV pipeline，还体现 RAG + LLM 应用能力。很多 AI 应用岗位关心模型结果如何解释、如何生成报告、如何防止幻觉，本项目正好覆盖这些点。

### 7.2 原始 LLM-based 思路：让 LLM 直接生成分析

最初想法是把 metrics 和 RAG context 交给 LLM，让它写 Engineering Interpretation。但这会带来风险：LLM 可能夸大、重复、编造 baseline，或者把中等指标写成优秀结果。

### 7.3 出现的问题：幻觉、过度夸大、编造指标

例如 PSNR 约 19.39、SSIM 约 0.699 时，LLM 曾倾向写成“near-perfect”或“excellent”。这与工程判断不符。

### 7.4 Rule-based 和 LLM-based 的区别

Rule-based 是 deterministic 的，输入同样 metrics 会得到同样结论。LLM-based 是语言生成，适合做解释补充，但不适合作为唯一结论来源。

### 7.5 为什么 rule-based 应该作为 source of truth

指标解释涉及工程责任。如果报告说“实时”或“优秀”，但 metrics 显示 {infer_ms} ms/frame 和中等 PSNR/SSIM，就会误导读者。因此主结论必须由规则生成。

### 7.6 为什么 LLM 只做 Additional Analysis 更安全

LLM Additional Analysis 可以解释趋势、提出下一步实验、连接项目背景，但不能修改指标表和规则分析。这种分层设计更安全，也更容易在面试中解释。

### 7.7 qwen2.5:1.5b 和 qwen2.5:3b 的差异

1.5B 通常更快，但更容易输出模板化或不稳定内容；3B 可能表达更好，但耗时更长，容易 timeout。

### 7.8 timeout 问题：为什么 3B 需要更长等待时间

本地 LLM 速度受硬件和模型大小影响。`--llm-timeout` 允许根据模型调整等待时间，而不是把超时写死。

### 7.9 sanitize 与 validation：防止 LLM 胡说

sanitize 先替换不安全正向夸大句，例如“near-perfect restoration”；validation 再检查是否编造数字或与 rule-based 冲突。安全否定句如“not near-perfect”会被允许。

### 7.10 LLM 调试文件的作用

- `llm_raw_output.md`：保存模型原始回答。
- `llm_sanitized_output.md`：保存替换不安全措辞后的版本。
- `llm_validation_log.txt`：记录命中的 phrase、句子、处理动作和最终结果。

## 8. 最终报告结构说明

### 8.1 Metric Summary：原始指标表

这是最基础的数据层，不能被 LLM 改写。所有解释都应回到这里验证。

### 8.2 Retrieved VSRNet Knowledge：RAG 检索到的项目知识

这部分来自 `documents/` 文档检索，提供项目背景、方法和评估语境。

### 8.3 Rule-Based Engineering Analysis：可信主分析

这是确定性主结论，解释 PSNR、SSIM、jitter、IoU continuity、confidence 和 runtime。

### 8.4 LLM Additional Analysis：补充解释层

只在开启 `--use-llm-report` 时出现。它可以帮助自然语言表达，但不应该被当成原始事实来源。

### 8.5 Recommendations / Warnings / Next Steps

Recommendations 给出工程建议，Warnings 说明当前限制，Next Steps 指明下一步实验或优化方向。

### 8.6 面试时如何解释“不完全相信 LLM 输出”

可以这样说：我把 LLM 放在补充层，而不是结论层。原始指标和 rule-based analysis 是 source of truth，LLM 输出必须经过 sanitize 和 validation 才能进入报告。

## 9. 面试高频问题准备

### 9.1 这个项目解决了什么问题？

它解决的是视频恢复模型的工程评估问题：不仅看 VSRNet 恢复后画面是否更好，还看下游 YOLO 检测是否更稳定，并把指标、项目知识和报告生成串成完整流程。

### 9.2 为什么 VSRNet 后面还要接 YOLO？

因为工程场景中恢复视频的价值通常体现在下游任务上。YOLO 可以帮助验证恢复结果是否改善或稳定了目标检测。

### 9.3 为什么要做 RAG？

RAG 让报告能引用项目本身的 summary、method notes 和 evaluation notes，减少报告脱离项目背景的风险。

### 9.4 为什么不用 LLM 直接写最终结论？

因为 LLM 容易夸大或编造。我的设计是 metrics 和 rule-based analysis 做可信主结论，LLM 只做 supplementary commentary。

### 9.5 你怎么防止 LLM 幻觉？

我做了三层：prompt 限制、sanitize 替换不安全措辞、validation 检查数字和结论冲突。同时保存 raw/sanitized/log 方便追踪。

### 9.6 rule-based 和 LLM 分析有什么区别？

rule-based 是确定性的、可复现的，适合作为 source of truth；LLM 更擅长语言组织和提出下一步思路，但必须受约束。

### 9.7 为什么要有 report-only / reuse 模式？

因为 VSRNet 推理耗时长，重复运行浪费时间和显存。reuse 模式可以复用 restored video、detection CSV 和 metrics JSON，只重新生成报告。

### 9.8 遇到过哪些工程 bug？怎么定位？

典型 bug 包括相对路径导致输出找不到、SSIM 通道维度错误、checkpoint 路径错误、LLM 过度夸大。我通常先读真实脚本接口，再打印绝对路径、cwd、输出文件列表和中间结果。

### 9.9 这个项目和传统 CV 项目相比有什么亮点？

它不是只输出恢复视频，而是形成 CV + downstream metrics + RAG + LLM report 的闭环，更接近真实 AI 应用系统。

### 9.10 如果继续优化，你会怎么做？

我会增加 baseline 对比、批量 vibration level 处理、更完整的 tracking 指标、可视化报告和更系统的测试。

## 10. 简历与面试表达模板

### 10.1 简历项目描述中文版

构建 VSRNet 工程评估助手，集成视频恢复、YOLO 下游检测、PSNR/SSIM 与检测稳定性指标计算、FAISS RAG 项目知识检索、确定性 rule-based 分析和带防幻觉校验的 LLM 补充评论，支持 report-only 复用模式以降低重复推理成本。

### 10.2 简历项目描述英文版

Built a CV + RAG + LLM engineering assistant for VSRNet, integrating video restoration, YOLO-based downstream evaluation, metric computation, FAISS retrieval, deterministic rule-based analysis, and validated LLM commentary with anti-hallucination safeguards.

### 10.3 面试 30 秒项目介绍

这个项目是一个围绕 VSRNet 的工程评估助手。我没有重写模型，而是把已有视频恢复代码包装成可运行 pipeline，后面接 YOLO 检测、指标计算、RAG 检索和报告生成。报告中 rule-based 分析是可信主结论，LLM 只做补充评论，并且有防幻觉校验。

### 10.4 面试 2 分钟项目介绍

我做的是 VSRNet Engineering Assistant，目标是评估振动退化视频恢复在工程上的实际效果。第一阶段先做 placeholder MVP，确保 restoration、detection、metrics、RAG 和 report 的接口跑通。第二阶段接入真实 VSRNet 推理和 YOLO 检测，计算 PSNR、SSIM、jitter、IoU continuity、confidence 和 inference time。后来我又加入 RAG，让报告能引用项目文档；加入 rule-based analysis，保证指标解释稳定可信；加入 LLM Additional Analysis，让本地 Ollama 或 OpenAI 只做补充评论，并用 sanitize + validation 防止幻觉。项目还支持 report-only 模式，避免每次改报告都重新跑几分钟的 VSRNet 推理。

### 10.5 技术亮点表达：CV + Metrics + RAG + LLM

技术亮点可以概括为：模型复用而非重写、下游任务评估、指标驱动解释、RAG 增强上下文、LLM 防幻觉补充评论、Windows 中文路径和复用模式工程化处理。

### 10.6 问到 bug 调试时的回答模板

我会先复现现象，然后定位是输入、路径、子进程 cwd、依赖版本还是输出验证的问题。例如 VSRNet 推理成功但主程序找不到输出时，我通过打印 cwd 和绝对输出路径发现是相对路径被解释到 VSRNet repo 下，于是统一改为绝对路径。

### 10.7 问到 LLM 幻觉时的回答模板

我没有让 LLM 直接决定结论，而是把它放在 Additional Analysis 层。核心结论来自 rule-based analysis。LLM 输出会先保存 raw，再 sanitize，最后 validation；如果编造数字或和规则结论冲突，就不会作为有效评论进入报告。

### 10.8 问到工程优化时的回答模板

我会优先做三件事：加入 baseline 对比，优化 VSRNet 推理速度，完善 tracking-based stability metrics。这样能让项目从单样本 demo 变成更系统的评估工具。

## 11. 后续优化方向

### 11.1 增加多 vibration level 自动批处理

让 low/mid/high 多组视频自动跑完整 pipeline，并生成对比报告。

### 11.2 增加 baseline 对比：FastDVDnet / VRT / Raw input

当前不能声称优于其他方法，因为缺少 baseline 指标。后续应加入 Raw、FastDVDnet、VRT 等对比。

### 11.3 增加更完整的 tracking-based jitter 计算

目前 jitter 基于检测框中心点，后续可以按稳定 track、目标类别和时间窗口细化。

### 11.4 增加更好的检测置信度分析

可以分析 confidence 分布、低置信帧、漏检帧，并结合可视化截图。

### 11.5 支持更好的可视化报告

增加曲线图、检测框连续性图、运行时间图和视频帧对比。

### 11.6 支持本地轻量 LLM 与 API LLM 双模式

保留 Ollama 本地隐私友好模式，同时支持 OpenAI API 生成更稳定的补充评论。

### 11.7 将项目整理成 GitHub README 与演示材料

后续可以准备 demo GIF、架构图、快速运行命令和面试讲解图。

## 面试时最值得强调的 5 个点

1. 这是一个完整 AI 工程系统，不是单点模型 demo。
2. 评估不仅看 PSNR/SSIM，还看 YOLO 下游稳定性。
3. 使用 RAG 把项目知识接入报告，减少脱离上下文的解释。
4. rule-based analysis 是可信主结论，LLM 只做补充并带防幻觉机制。
5. 通过 reuse/report-only 模式解决昂贵推理反复运行的问题。

## 不要过度声称的内容

- 不要说当前恢复质量“excellent”或“near-perfect”，因为 PSNR={psnr}、SSIM={ssim} 只能支持中等恢复质量。
- 不要说当前系统“real-time”，因为 inference_time_ms 约 {infer_ms} ms/frame。
- 不要说已经优于 FastDVDnet、VRT 或传统方法，因为当前没有 baseline 对比指标。
- 不要说 YOLO 检测非常可靠，因为 confidence_mean 约 {conf_mean} 偏低。
- 不要把 LLM Additional Analysis 当作事实来源，最终仍以 raw metrics 和 rule-based analysis 为准。
"""


def set_cell_shading(cell, fill: str):
    """Set table cell shading."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    """Set cell text with consistent formatting."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a simple full-width table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], str(text))
    doc.add_paragraph()


def set_update_fields_on_open(doc: Document) -> None:
    """Ask Word to refresh fields such as the table of contents on open."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def extract_toc_entries(markdown: str) -> list[tuple[int, str]]:
    """Collect headings used to seed the DOCX table of contents."""
    entries: list[tuple[int, str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            entries.append((1, line[3:].strip()))
        elif line.startswith("### "):
            entries.append((2, line[4:].strip()))
    return entries


def add_table_of_contents(doc: Document, entries: list[tuple[int, str]]) -> None:
    """Insert a Word table-of-contents field with a visible seeded result."""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("目录")
    heading_run.bold = True
    heading_run.font.size = Pt(16)

    paragraph = doc.add_paragraph()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph.add_run()._r.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    paragraph.add_run()._r.append(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    result_paragraph = paragraph
    if entries:
        for level, title in entries:
            result_paragraph = doc.add_paragraph()
            result_paragraph.paragraph_format.left_indent = Inches(0.24 * (level - 1))
            result_paragraph.paragraph_format.space_after = Pt(2)
            run = result_paragraph.add_run(title)
            run.font.size = Pt(10.5)
    else:
        paragraph.add_run("打开文档后目录会自动更新；如未更新，可在 Word 中右键目录选择“更新域”。")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    result_paragraph.add_run()._r.append(end)

    doc.add_page_break()


def add_markdown_to_docx(markdown: str, docx_path: Path):
    """Create a readable DOCX from the generated Markdown."""
    doc = Document()
    set_update_fields_on_open(doc)
    toc_entries = extract_toc_entries(markdown)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "VSRNet Engineering Assistant 调试与面试复盘"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)

    in_code = False
    table_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        lines = table_buffer
        table_buffer = []
        if len(lines) < 2:
            for line in lines:
                doc.add_paragraph(line)
            return
        headers = [cell.strip(" `") for cell in lines[0].strip().strip("|").split("|")]
        rows = []
        for line in lines[2:]:
            cells = [cell.strip(" `") for cell in line.strip().strip("|").split("|")]
            if len(cells) == len(headers):
                rows.append(cells)
        if rows:
            add_table(doc, headers, rows)

    toc_inserted = False
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            paragraph = doc.add_paragraph(style="Normal")
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(line[2:], style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not toc_inserted:
                add_table_of_contents(doc, toc_entries)
                toc_inserted = True
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:3].replace(".", "").isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line[3:], style="List Number")
        else:
            doc.add_paragraph(line)

    flush_table()
    doc.save(docx_path)


def main():
    """Write Markdown and DOCX documents."""
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_metrics()
    markdown = build_markdown(metrics)
    MD_PATH.write_text(markdown, encoding="utf-8")
    add_markdown_to_docx(markdown, DOCX_PATH)
    print(f"Wrote Markdown: {MD_PATH}")
    print(f"Wrote DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
